"""
file_utils.py - AI-Powered Document Parsing
Uses Gemini multimodal for PDF/DOCX extraction with local fallback.
"""

import base64
import os
from pathlib import Path

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
import pdfplumber
from docx import Document as DocxDocument

from .llm import get_llm, get_fallback_chain, get_manual_review_response

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
MAX_FILE_SIZE_MB = MAX_FILE_SIZE_BYTES / (1024 * 1024)


class FileProcessingError(Exception):
    """Base exception for file processing errors."""
    pass


class FileTooLargeError(FileProcessingError):
    """Raised when file exceeds maximum size limit."""
    def __init__(self, file_path: str, size_mb: float):
        self.file_path = file_path
        self.size_mb = size_mb
        super().__init__(
            f"File too large: {file_path} ({size_mb:.1f} MB). "
            f"Maximum allowed size is {MAX_FILE_SIZE_MB:.0f} MB."
        )


class UnsupportedFormatError(FileProcessingError):
    """Raised when file format is not supported."""
    pass


class ExtractionError(FileProcessingError):
    """Raised when text extraction fails."""
    pass



# File Validation


def get_file_size_mb(file_path: str) -> float:
    """Get file size in megabytes."""
    size_bytes = os.path.getsize(file_path)
    return size_bytes / (1024 * 1024)


def validate_file_size(file_path: str) -> None:
    """
    Validate that file size is within acceptable limits.
    
    Args:
        file_path: Path to the file
        
    Raises:
        FileTooLargeError: If file exceeds maximum size
    """
    size_bytes = os.path.getsize(file_path)
    size_mb = size_bytes / (1024 * 1024)
    
    if size_bytes > MAX_FILE_SIZE_BYTES:
        raise FileTooLargeError(file_path, size_mb)



# File Reading Utilities


def get_mime_type(file_path: str) -> str:
    """
    Determine the MIME type based on file extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        MIME type string
    """
    extension = Path(file_path).suffix.lower()
    
    mime_types = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
    }
    
    return mime_types.get(extension, "application/octet-stream")


def encode_file_to_base64(file_path: str) -> str:
    """
    Read a file and encode it to base64 string.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Base64 encoded string of the file contents
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file cannot be read
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        
        if not file_bytes:
            raise ValueError(f"File is empty: {file_path}")
        
        return base64.b64encode(file_bytes).decode("utf-8")
    
    except Exception as e:
        if isinstance(e, (FileNotFoundError, ValueError)):
            raise
        raise ValueError(f"Failed to read file '{file_path}': {str(e)}")


def validate_resume_file(file_path: str) -> dict:
    """
    Validate that the resume file exists, has a supported format, and is within size limits.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        dict with file info: {"size_mb": float, "format": str}
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        UnsupportedFormatError: If the file format is not supported
        FileTooLargeError: If the file exceeds size limit
    """
    path = Path(file_path)
    
    # Check file exists
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    # Check file format
    supported_extensions = [".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"]
    if path.suffix.lower() not in supported_extensions:
        raise UnsupportedFormatError(
            f"Unsupported file format: {path.suffix}\n"
            f"Supported formats: {', '.join(supported_extensions)}"
        )
    
    # Check file size
    size_mb = get_file_size_mb(file_path)
    validate_file_size(file_path)
    
    return {
        "size_mb": size_mb,
        "format": path.suffix.lower(),
    }



# Local File Extraction (Fallback when Gemini fails)


def extract_text_from_pdf_local(file_path: str) -> str:
    """
    Extract text from PDF using pdfplumber (local, no API needed).
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from all pages
        
    Raises:
        ExtractionError: If extraction fails
    """
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                raise ExtractionError("PDF file has no pages - may be corrupted")
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_error:
                    # Log but continue - some pages might be image-only
                    print(f"      ⚠️ Could not extract text from page {page_num + 1}")
        
        if not text_parts:
            raise ExtractionError(
                "No text could be extracted from PDF. "
                "The file may be image-only, password-protected, or corrupted."
            )
        
        return "\n\n".join(text_parts)
    except ExtractionError:
        raise
    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            raise ExtractionError("PDF is password-protected. Please provide an unlocked file.")
        elif "invalid" in error_msg or "corrupted" in error_msg:
            raise ExtractionError("PDF file is corrupted or invalid.")
        raise ExtractionError(f"PDF extraction failed: {str(e)[:100]}")


def extract_text_from_docx_local(file_path: str) -> str:
    """
    Extract text from DOCX using python-docx (local, no API needed).
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from all paragraphs and tables
        
    Raises:
        ExtractionError: If extraction fails
    """
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        if not text_parts:
            raise ExtractionError(
                "No text could be extracted from DOCX. "
                "The file may be empty or contain only images."
            )
        
        return "\n\n".join(text_parts)
    except ExtractionError:
        raise
    except Exception as e:
        error_msg = str(e).lower()
        if "not a valid" in error_msg or "package" in error_msg:
            raise ExtractionError("DOCX file is corrupted or not a valid Word document.")
        raise ExtractionError(f"DOCX extraction failed: {str(e)[:100]}")


def extract_text_locally(file_path: str) -> str:
    """
    Extract text from a file using local libraries (no API).
    
    This is the fallback when AI-based extraction fails.
    Supports PDF and DOCX files.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text
        
    Raises:
        ExtractionError: If extraction fails or format not supported
    """
    extension = Path(file_path).suffix.lower()
    
    if extension == ".pdf":
        return extract_text_from_pdf_local(file_path)
    elif extension in [".docx", ".doc"]:
        return extract_text_from_docx_local(file_path)
    else:
        raise ExtractionError(
            f"Local extraction not supported for {extension}. "
            f"Only PDF and DOCX are supported for offline fallback."
        )


def clean_text_with_groq(raw_text: str) -> str:
    """
    Use Groq/Llama to clean and format extracted text.
    
    This is optional - if Groq fails, we just return the raw text.
    
    Args:
        raw_text: Raw text extracted from document
        
    Returns:
        Cleaned text (or original if Groq fails)
    """
    try:
        from llm import get_llm
        
        # Try Groq for text cleaning
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            print("   ℹ️  No GROQ_API_KEY set, using raw extracted text")
            return raw_text
        
        llm = get_llm(model="llama-3.3-70b-versatile", temperature=0.0, provider="groq")
        
        prompt = f"""Clean and format the following resume text. 
Remove any artifacts, fix formatting issues, but preserve all content exactly.
Do NOT summarize or remove any information.

TEXT:
{raw_text}

CLEANED TEXT:"""
        
        response = llm.invoke(prompt)
        cleaned = response.content.strip()
        
        if cleaned and len(cleaned) > len(raw_text) * 0.5:
            return cleaned
        return raw_text
        
    except Exception as e:
        print(f"   ℹ️  Groq cleaning skipped: {str(e)[:50]}")
        return raw_text



# AI-Powered File Parser Agent


FILE_PARSER_PROMPT = """You are a document text extraction specialist.

Your task is to extract ALL text content from the provided document (resume/CV).

INSTRUCTIONS:
1. Read the entire document carefully
2. Extract ALL text content, preserving the logical structure
3. Include text from headers, sections, bullet points, tables, and any other elements
4. Maintain the order of information as it appears in the document
5. If there are multiple columns, read left to right, top to bottom
6. Include dates, company names, job titles, skills, education details - everything

OUTPUT:
Return ONLY the extracted text content. Do not add any commentary or formatting.
Just output the raw text from the document, with line breaks to preserve structure.

Extract the text now:"""


def extract_text_with_ai(file_path: str, max_retries: int = 2) -> str:
    """
    Extract text from a document using Gemini's multimodal capabilities.
    
    This is the FILE PARSER AGENT - it uses AI to read and extract text
    from PDF, DOCX, or image files directly.
    
    Features:
    - File size validation
    - Multi-model fallback (tries different Gemini models)
    - Retry logic for transient failures
    
    Note: Only Gemini supports multimodal (PDF reading). Groq/Llama cannot read PDFs.
    
    Args:
        file_path: Path to the resume file (PDF, DOCX, or image)
        max_retries: Maximum retry attempts per model
        
    Returns:
        Extracted text content from the document
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        FileTooLargeError: If file exceeds size limit
        UnsupportedFormatError: If format not supported
        ExtractionError: If text extraction fails after all retries
    """
    import time
    
    # Validate file (checks existence, format, and size)
    file_info = validate_resume_file(file_path)
    print(f"   📁 File validated: {file_info['format']} ({file_info['size_mb']:.1f} MB)")
    
    # Get file details
    mime_type = get_mime_type(file_path)
    file_data = encode_file_to_base64(file_path)
    
    # Gemini models to try (only Gemini supports multimodal PDF reading)
    gemini_models = [
        "gemini-2.5-flash",
        "gemini-2.0-flash", 
        "gemini-2.0-flash-lite",
    ]
    
    last_error = None
    
    for model_idx, model in enumerate(gemini_models):
        for attempt in range(max_retries):
            try:
                # Get LLM for this model
                llm = get_llm(model=model, temperature=0.0)
                
                attempt_info = f"Model {model_idx + 1}/{len(gemini_models)}: {model}"
                if attempt > 0:
                    attempt_info += f" (retry {attempt + 1}/{max_retries})"
                print(f"   🔄 {attempt_info}")
                
                # Create multimodal message with file attachment
                message = HumanMessage(
                    content=[
                        {
                            "type": "text",
                            "text": FILE_PARSER_PROMPT,
                        },
                        {
                            "type": "media",
                            "mime_type": mime_type,
                            "data": file_data,
                        },
                    ]
                )
                
                # Invoke the LLM with the file
                response = llm.invoke([message])
                extracted_text = response.content.strip()
                
                if not extracted_text:
                    raise ExtractionError("No text could be extracted from the document")
                
                print(f"   ✓ Text extracted successfully with {model}")
                return extracted_text
                
            except Exception as e:
                last_error = str(e)
                error_lower = last_error.lower()
                
                if "429" in last_error or "rate" in error_lower or "quota" in error_lower:
                    # Rate limit - wait briefly then try next model
                    print(f"   ⏳ {model} rate limited, trying next model...")
                    time.sleep(2)
                    break  # Move to next model
                elif "404" in last_error or "not found" in error_lower:
                    # Model not available - try next
                    print(f"   ⚠️ {model} not available, trying next...")
                    break  # Move to next model
                elif "401" in last_error or "403" in last_error:
                    # Auth error - no point continuing
                    raise ExtractionError(f"API authentication failed: {last_error[:100]}")
                else:
                    # Other error - retry with same model
                    print(f"   ⚠️ Error with {model}: {last_error[:60]}...")
                    time.sleep(1)
    
    # All Gemini models failed - try local extraction as fallback
    print("   🔄 All Gemini models exhausted. Trying local extraction...")
    
    extension = Path(file_path).suffix.lower()
    if extension in [".pdf", ".docx", ".doc"]:
        try:
            print(f"   📥 Extracting {extension} locally with pdfplumber/python-docx...")
            raw_text = extract_text_locally(file_path)
            print(f"   ✓ Local extraction successful ({len(raw_text)} characters)")
            
            # Optionally clean with Groq
            print("   🔧 Attempting to clean text with Groq...")
            cleaned_text = clean_text_with_groq(raw_text)
            
            if cleaned_text != raw_text:
                print("   ✓ Text cleaned with Groq")
            
            return cleaned_text
            
        except ExtractionError as e:
            raise ExtractionError(
                f"Both AI and local extraction failed. "
                f"AI error: {last_error[:100] if last_error else 'Unknown'}. "
                f"Local error: {str(e)}"
            )
    else:
        # Images can't be extracted locally
        raise ExtractionError(
            f"Failed to extract text from document. All Gemini models exhausted. "
            f"Local extraction not supported for {extension} files. "
            f"Last error: {last_error[:200] if last_error else 'Unknown'}"
        )



# Job Description Reading (Plain Text)


def read_job_description(file_path: str) -> str:
    """
    Read job description from a text file.
    
    Args:
        file_path: Path to the job description text file
        
    Returns:
        Job description text
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is empty
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Job description file not found: {file_path}")
    
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        
        if not content:
            raise ValueError(f"Job description file is empty: {file_path}")
        
        return content
    
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, "r", encoding="latin-1") as f:
            content = f.read().strip()
        
        if not content:
            raise ValueError(f"Job description file is empty: {file_path}")
        
        return content



# Convenience Function (Backwards Compatible)


def extract_text(file_path: str) -> str:
    """
    Extract text from a resume file using AI.
    
    This is an alias for extract_text_with_ai() to maintain backwards
    compatibility with the existing codebase.
    
    Args:
        file_path: Path to the resume file (PDF, DOCX, or image)
        
    Returns:
        Extracted text content
    """
    return extract_text_with_ai(file_path)
